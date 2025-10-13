"""DOCX document parser."""
import io
from typing import Dict, List, Any
from docx import Document

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../'))

from shared.logger import setup_logger

logger = setup_logger(__name__)


class DOCXParser:
    """Parse DOCX documents to extract text and structure."""
    
    @staticmethod
    def extract_text(file_bytes: bytes) -> str:
        """Extract text from DOCX.
        
        Args:
            file_bytes: DOCX file bytes
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraphs)
            logger.info(f"Extracted {len(paragraphs)} paragraphs from DOCX")
            return text
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {str(e)}")
            raise
    
    @staticmethod
    def extract_tables(file_bytes: bytes) -> List[List[List[str]]]:
        """Extract tables from DOCX.
        
        Args:
            file_bytes: DOCX file bytes
            
        Returns:
            List of tables
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            all_tables = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                all_tables.append(table_data)
            
            logger.info(f"Extracted {len(all_tables)} table(s) from DOCX")
            return all_tables
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX tables: {str(e)}")
            return []
    
    @staticmethod
    def extract_metadata(file_bytes: bytes) -> Dict[str, Any]:
        """Extract DOCX metadata.
        
        Args:
            file_bytes: DOCX file bytes
            
        Returns:
            Document metadata
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            core_props = doc.core_properties
            
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0,
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX metadata: {str(e)}")
            return {}
    
    @staticmethod
    def parse_structured(file_bytes: bytes) -> Dict[str, Any]:
        """Parse DOCX with structure (headings, paragraphs, tables).
        
        Args:
            file_bytes: DOCX file bytes
            
        Returns:
            Structured document data
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            structured_data = {
                'sections': [],
                'tables': [],
                'metadata': DOCXParser.extract_metadata(file_bytes)
            }
            
            current_section = None
            
            for element in doc.element.body:
                # Check if it's a paragraph
                if element.tag.endswith('p'):
                    for para in doc.paragraphs:
                        if para._element == element:
                            # Check if it's a heading
                            if para.style.name.startswith('Heading'):
                                current_section = {
                                    'heading': para.text,
                                    'level': para.style.name,
                                    'content': []
                                }
                                structured_data['sections'].append(current_section)
                            elif current_section:
                                current_section['content'].append(para.text)
                            else:
                                # No section yet, create default
                                if not structured_data['sections']:
                                    structured_data['sections'].append({
                                        'heading': 'Document',
                                        'level': 'default',
                                        'content': []
                                    })
                                    current_section = structured_data['sections'][0]
                                current_section['content'].append(para.text)
                
                # Check if it's a table
                elif element.tag.endswith('tbl'):
                    for table in doc.tables:
                        if table._element == element:
                            table_data = []
                            for row in table.rows:
                                row_data = [cell.text.strip() for cell in row.cells]
                                table_data.append(row_data)
                            structured_data['tables'].append(table_data)
            
            return structured_data
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX structure: {str(e)}")
            return {'sections': [], 'tables': [], 'metadata': {}}

